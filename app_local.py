import sys
import time
import threading
from pathlib import Path

PROJECT_DIR = Path(__file__).parent
REPO_DIR = PROJECT_DIR / "privacy-filter"
sys.path.insert(0, str(REPO_DIR))

def _patch_gradio_client():
    try:
        from gradio_client import utils as _cu
        _orig_get_type = _cu.get_type
        if getattr(_orig_get_type, "_patched", False):
            return
        def _safe_get_type(schema):
            if not isinstance(schema, dict):
                return "Any"
            return _orig_get_type(schema)
        _safe_get_type._patched = True
        _cu.get_type = _safe_get_type
        _orig_json = _cu._json_schema_to_python_type
        def _safe_json(schema, defs):
            if not isinstance(schema, dict):
                return "Any"
            return _orig_json(schema, defs)
        _cu._json_schema_to_python_type = _safe_json
    except Exception:
        pass

_patch_gradio_client()

import gradio as gr

_model = None


def _checkpoint_dir() -> Path:
    """Return the default OPF checkpoint directory."""
    return Path.home() / ".opf" / "privacy_filter"


def _is_partial_checkpoint(model_dir: Path) -> bool:
    """Return True if ``model_dir`` exists but is missing required files."""
    if not model_dir.exists() or not model_dir.is_dir():
        return False
    if not (model_dir / "config.json").is_file():
        return True
    return not any(model_dir.glob("*.safetensors"))


def _ensure_model_present(progress_callback=None) -> None:
    """Download the OPF checkpoint if it is missing or in a partial state."""
    from model_update import download_model_update
    success, message = download_model_update(progress_callback=progress_callback)
    if not success:
        raise RuntimeError(
            f"Could not download the PII model: {message}. "
            f"Check your internet connection or remove the partial "
            f"checkpoint at {_checkpoint_dir()} and try again."
        )


def get_model():
    global _model
    if _model is not None:
        return _model
    print("Loading Privacy Filter model...")

    from opf._api import OPF

    try:
        _model = OPF(device="cpu")
        print("[OK] Model loaded")
        return _model
    except (FileNotFoundError, RuntimeError) as exc:
        # OPF raises FileNotFoundError when the directory is missing
        # required files (e.g. config.json), and RuntimeError for other
        # validation failures. In both cases, a partial or missing
        # checkpoint is the most common cause -- try to recover by
        # downloading a fresh copy.
        message = str(exc)
        recoverable = (
            "missing config.json" in message
            or "no .safetensors" in message
            or "Checkpoint directory not found" in message
        )
        if not recoverable or not _is_partial_checkpoint(_checkpoint_dir()):
            print(f"[ERROR] {exc}")
            raise

        print("[INFO] Checkpoint is missing or partial. Downloading...")
        _ensure_model_present(
            progress_callback=lambda msg, pct: print(f"[INFO] {msg}")
        )
        _model = OPF(device="cpu")
        print("[OK] Model loaded after download")
        return _model
    except Exception as e:
        print(f"[ERROR] {e}")
        raise

def extract_text_from_pdf(pdf_path):
    try:
        import fitz
        doc = fitz.open(pdf_path)
        try:
            pages = [page.get_text() for page in doc]
        finally:
            doc.close()
        return "\n".join(pages)
    except Exception as e:
        print(f"[PDF ERROR] {e}")
    return None

def extract_text_from_docx(docx_path):
    try:
        from docx import Document
        doc = Document(docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text
    except Exception as e:
        print(f"[DOCX ERROR] {e}")
    return None

def redact_pdf(input_path, detected_spans):
    import os
    import uuid
    import fitz
    out_path = os.path.join(
        os.environ.get("TEMP", os.path.dirname(input_path)),
        f"redacted_{uuid.uuid4().hex}.pdf",
    )
    doc = fitz.open(input_path)
    for page in doc:
        for span in detected_spans:
            old_text = span.text
            new_text = span.placeholder
            if not old_text or old_text == new_text:
                continue
            results = page.search_for(old_text)
            for rect in results:
                page.add_redact_annot(
                    rect, text=new_text, fontsize=9,
                    fontname="helv", fill=(1, 1, 1), text_color=(0, 0, 0),
                    align=0
                )
        page.apply_redactions()
    doc.save(out_path)
    doc.close()
    return out_path

def redact_docx(input_path, detected_spans):
    """Apply PII redactions to a DOCX while preserving run-level formatting.

    Replaces ``old_text`` with ``new_text`` inside each affected paragraph by
    splitting the original runs at match boundaries, so the surrounding runs
    keep their original fonts, sizes, bold/italic flags, etc.
    """
    import os
    import uuid
    from docx import Document
    out_path = os.path.join(
        os.environ.get("TEMP", os.path.dirname(input_path)),
        f"redacted_{uuid.uuid4().hex}.docx",
    )
    doc = Document(input_path)
    for span in detected_spans:
        old_text = span.text
        new_text = span.placeholder
        if not old_text or old_text == new_text:
            continue
        for para in doc.paragraphs:
            if old_text not in para.text:
                continue
            _replace_text_in_paragraph(para, old_text, new_text)
    doc.save(out_path)
    return out_path


def _replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace ``old_text`` with ``new_text`` in ``paragraph`` run-by-run.

    Splits runs at the first match boundary so the resulting placeholder sits
    in its own run; the run-level formatting of the surrounding text is
    preserved.
    """
    full_text = paragraph.text
    start = full_text.find(old_text)
    if start < 0:
        return
    end = start + len(old_text)

    # Build the new sequence of (text, run_template) for the paragraph.
    # Each overlapping run produces a leading slice (if any) and a trailing
    # slice (if any); the placeholder is inserted exactly once for the whole
    # match even when it spans multiple runs.
    segments: list[tuple[str, object]] = []
    cursor = 0
    placeholder_inserted = False
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        run_start = cursor
        run_end = cursor + len(run_text)
        if run_end <= start or run_start >= end:
            segments.append((run_text, run))
        else:
            if run_start < start:
                segments.append((run_text[: start - run_start], run))
            if not placeholder_inserted:
                segments.append((new_text, None))
                placeholder_inserted = True
            if run_end > end:
                segments.append((run_text[end - run_start :], run))
        cursor = run_end

    if not segments:
        return

    # Find the first run with content to use as the formatting template for
    # the placeholder.
    template_for_placeholder = None
    for run in paragraph.runs:
        if run.text:
            template_for_placeholder = run
            break

    # Clear all original runs.
    for run in paragraph.runs:
        run.text = ""

    # Reuse the first run for the first segment, then add new runs for the rest.
    first_text, first_template = segments[0]
    paragraph.runs[0].text = first_text
    if first_template is not None:
        _copy_font(first_template.font, paragraph.runs[0].font)

    for text, template in segments[1:]:
        new_run = paragraph.add_run(text)
        if template is not None:
            _copy_font(template.font, new_run.font)
        elif template_for_placeholder is not None:
            _copy_font(template_for_placeholder.font, new_run.font)


def _copy_font(src_font, dst_font):
    """Copy font attributes from ``src_font`` to ``dst_font`` if set."""
    if src_font.name:
        dst_font.name = src_font.name
    if src_font.size:
        dst_font.size = src_font.size
    if src_font.bold is not None:
        dst_font.bold = src_font.bold
    if src_font.italic is not None:
        dst_font.italic = src_font.italic

def redact_text(text):
    if not text or not text.strip():
        return "", "Enter some text."
    try:
        model = get_model()
        start = time.time()
        result = model.redact(text)
        elapsed = time.time() - start
        redacted = result.redacted_text if hasattr(result, 'redacted_text') else str(result)
        spans = result.detected_spans if hasattr(result, 'detected_spans') else []
        if spans:
            lines = [f"**{len(spans)} entities detected** ({elapsed:.1f}s)", ""]
            lines.extend(
                f"- `{s.label if hasattr(s, 'label') else '?'}`: {s.text if hasattr(s, 'text') else ''}"
                for s in spans
            )
            summary = "\n".join(lines)
        else:
            summary = f"_No PII entities detected_ ({elapsed:.1f}s)"
        return redacted, summary
    except Exception as e:
        return text, f"Error: {e}"

def redact_file(file, progress=gr.Progress()):
    if file is None:
        return "_Upload a file._", None
    try:
        progress((0, 5), desc="Loading model...")
        model = get_model()

        progress((1, 5), desc="Reading file...")
        path = Path(file.name)
        ext = path.suffix.lower()
        if ext == ".pdf":
            text = extract_text_from_pdf(str(path))
            if text is None:
                return "**Error** reading PDF. Install: `pip install PyMuPDF`", None
        elif ext == ".docx":
            text = extract_text_from_docx(str(path))
            if text is None:
                return "**Error** reading DOCX. Check the console.", None
        else:
            text = path.read_text(encoding="utf-8", errors="replace")

        progress((2, 5), desc="Detecting PII...")
        start = time.time()
        result = model.redact(text)
        elapsed = time.time() - start
        spans = result.detected_spans if hasattr(result, 'detected_spans') else []

        progress((3, 5), desc="Generating redacted output...")
        legend_parts = [
            "### Redaction Result",
            "",
            f"Processed in **{elapsed:.1f}s** -- **{len(spans)}** entities detected",
            "",
        ]
        if spans:
            legend_parts.append("| # | Type | Original | Replacement |")
            legend_parts.append("|--:|------|----------|------------|")
            for i, s in enumerate(spans, 1):
                label = s.label if hasattr(s, "label") else "?"
                txt = s.text if hasattr(s, "text") else ""
                ph = s.placeholder if hasattr(s, "placeholder") else ""
                legend_parts.append(f"| {i} | `{label}` | {txt} | {ph} |")
        else:
            legend_parts.append("_No PII entities detected._")
        legend = "\n".join(legend_parts)

        progress((4, 5), desc="Creating output file...")
        out_path = None
        if ext == ".pdf" and spans:
            out_path = redact_pdf(str(path), spans)
        elif ext == ".docx" and spans:
            out_path = redact_docx(str(path), spans)
        progress((5, 5), desc="Done")
        return legend, out_path
    except Exception as e:
        return f"**Error:** {e}", None


# ============================================================
#  APP UPDATE FUNCTIONS
# ============================================================

_app_update_info = None


def _check_app_update_background():
    """Background thread that checks for app updates."""
    global _app_update_info
    try:
        from app_update import check_for_app_update
        _app_update_info = check_for_app_update()
    except Exception as exc:
        print(f"[UPDATE] App update check failed: {exc}")
        return
    if _app_update_info.update_available:
        print(
            f"[UPDATE] App update available: "
            f"v{_app_update_info.current_version} -> v{_app_update_info.latest_version}"
        )


# ============================================================
#  MODEL UPDATE FUNCTIONS
# ============================================================

_model_update_info = None


def _check_model_update_background():
    """Background thread that checks for model updates."""
    global _model_update_info
    try:
        from model_update import check_for_model_update
        _model_update_info = check_for_model_update()
    except Exception as exc:
        print(f"[UPDATE] Model update check failed: {exc}")
        return
    if _model_update_info.update_available:
        print(
            f"[UPDATE] Model update available: "
            f"{_model_update_info.current_date} -> {_model_update_info.latest_date}"
        )

def install_app_update(progress=gr.Progress()):
    """Download and install the app update."""
    global _app_update_info
    
    if not _app_update_info or not _app_update_info.update_available:
        return "_No update available._"
    
    try:
        progress((0.1, 1.0), desc="Updating application...")
        
        # Use git pull if no download URL (ZIP not attached to release)
        if not _app_update_info.download_url:
            import subprocess
            project_dir = Path(__file__).parent
            
            progress((0.3, 1.0), desc="Pulling latest code...")
            result = subprocess.run(
                ["git", "pull", "origin", "main"],
                cwd=str(project_dir),
                capture_output=True,
                text=True,
                timeout=60,
            )
            
            if result.returncode != 0:
                return f"**Git pull failed:** {result.stderr}\n\nPlease update manually:\n```\ncd {project_dir}\ngit pull\n.venv\\Scripts\\pip.exe install -e .\\privacy-filter\n```"
            
            progress((0.7, 1.0), desc="Installing dependencies...")
            venv_pip = project_dir / ".venv" / "Scripts" / "pip.exe"
            if venv_pip.exists():
                subprocess.run(
                    [str(venv_pip), "install", "-e", str(project_dir / "privacy-filter")],
                    cwd=str(project_dir),
                    capture_output=True,
                    timeout=120,
                )
            
            progress((1.0, 1.0), desc="Done!")
            from app_update import restart_app
            threading.Timer(2.0, restart_app).start()
            return f"**Updated to v{_app_update_info.latest_version}**\n\nRestarting in 2 seconds..."
        
        # Download ZIP if available
        from app_update import download_and_install_update, restart_app
        
        def update_progress(message, pct):
            progress((pct, 1.0), desc=message)
        
        success, message = download_and_install_update(
            _app_update_info.download_url,
            progress_callback=update_progress,
        )
        
        if success:
            progress((1.0, 1.0), desc="Done!")
            threading.Timer(2.0, restart_app).start()
            return f"**{message}**\n\nRestarting in 2 seconds..."
        else:
            return f"**Update failed:** {message}\n\nPlease update manually:\n```\ncd {Path(__file__).parent}\ngit pull\n.venv\\Scripts\\pip.exe install -e .\\privacy-filter\n```"
    
    except Exception as e:
        return f"**Error:** {e}"


def install_model_update(progress=gr.Progress()):
    """Download and install the model update."""
    global _model_update_info

    if not _model_update_info or not _model_update_info.update_available:
        return "_No model update available._"

    try:
        from model_update import download_model_update

        def update_progress(message, pct):
            progress((pct, 1.0), desc=message)

        success, message = download_model_update(
            progress_callback=update_progress,
        )

        if success:
            # Reload the model
            global _model
            _model = None
            get_model()

            progress((1.0, 1.0), desc="Done!")
            return f"**{message}**"
        else:
            return f"**Update failed:** {message}"

    except Exception as e:
        return f"**Error:** {e}"


def _hide_three():
    """Return three Gradio updates that hide their respective widgets."""
    return (gr.update(visible=False),) * 3


def create_ui():
    with gr.Blocks(title="Privacy Filter - Local") as app:
        # Read current version
        try:
            from app_update import get_local_version
            current_version = get_local_version()
        except Exception:
            current_version = "unknown"
        
        # Update banner (initially hidden)
        with gr.Row():
            update_banner = gr.Markdown(
                value="",
                visible=False,
            )
        
        with gr.Row():
            update_btn = gr.Button(
                "Update now",
                variant="primary",
                visible=False,
                scale=1,
            )
            later_btn = gr.Button(
                "Later",
                variant="secondary",
                visible=False,
                scale=1,
            )
        
        update_msg = gr.Markdown()

        update_btn.click(
            fn=install_app_update,
            outputs=update_msg,
        )

        later_btn.click(
            fn=_hide_three,
            outputs=[update_banner, update_btn, later_btn],
        )

        # Model update banner (initially hidden)
        with gr.Row():
            model_update_banner = gr.Markdown(
                value="",
                visible=False,
            )

        with gr.Row():
            model_update_btn = gr.Button(
                "Update model",
                variant="primary",
                visible=False,
                scale=1,
            )
            model_later_btn = gr.Button(
                "Later",
                variant="secondary",
                visible=False,
                scale=1,
            )

        model_update_msg = gr.Markdown()

        model_update_btn.click(
            fn=install_model_update,
            outputs=model_update_msg,
        )

        model_later_btn.click(
            fn=_hide_three,
            outputs=[model_update_banner, model_update_btn, model_later_btn],
        )

        gr.Markdown(f"# Privacy Filter - Local\n*100% local PII detection* | v{current_version}")

        with gr.Tab("Text"):
            with gr.Row():
                inp = gr.Textbox(
                    label="Text to analyze",
                    lines=5,
                    placeholder="My name is John, email: john@example.com, phone: +1 555 123 4567"
                )
                out = gr.Textbox(label="Redacted output", lines=5)
            btn = gr.Button("Detect PII", variant="primary")
            info = gr.Markdown("_Enter text and click Detect._")
            btn.click(fn=redact_text, inputs=inp, outputs=[out, info])
            gr.Examples(
                examples=[
                    "Hi, I'm John Smith. My email is john.smith@example.com and my SSN is 123-45-6789.",
                    "Call me at +1 555 987 6543 or email support@company.org",
                    "The meeting is on 03/15/2026. Account: 4532-1234-5678-9012",
                ],
                inputs=inp
            )

        with gr.Tab("Files"):
            gr.Markdown("Upload text or PDF files to redact PII.")
            finp = gr.File(
                label="Upload file",
                file_types=[".txt",".md",".csv",".json",".log",".py",".js",".xml",".html",".pdf",".docx"]
            )
            fbtn = gr.Button("Process File", variant="primary")
            flegend = gr.Markdown()
            fpdf = gr.File(label="Redacted file (PDF/DOCX)", visible=True)
            fbtn.click(fn=redact_file, inputs=finp, outputs=[flegend, fpdf])

        with gr.Tab("Info"):
            gr.Markdown("""
            ## PII Categories
            
            | Category | Description |
            |----------|-------------|
            | PERSON | Person names |
            | EMAIL | Email addresses |
            | PHONE | Phone numbers |
            | ADDRESS | Postal addresses |
            | DATE | Personal dates |
            | URL | Web links |
            | ACCOUNT_NUMBER | Bank accounts, cards |
            | SECRET | Passwords, API keys |
            
            ## Supported Formats
            
            - Text: .txt, .md, .csv, .json, .log, .py, .js, .xml, .html
            - PDF: .pdf (returns redacted PDF)
            - DOCX: .docx (returns redacted DOCX)
            
            ## Security
            
            - 100% local - nothing is sent to the internet
            - Model runs on your PC
            - Apache 2.0 license
            """)
            manual_update_btn = gr.Button("Update model", variant="secondary")
            manual_update_msg = gr.Markdown()
            manual_update_btn.click(fn=install_model_update, outputs=manual_update_msg)

        # Load update check on app start
        def check_and_update_banner():
            """Wait for update checks, then return banner widget updates."""
            _wait_for_update_info()

            return (
                *_build_app_banner_updates(),
                *_build_model_banner_updates(),
            )

        def _wait_for_update_info(timeout_seconds: float = 5.0) -> None:
            """Block until both background update checks have completed."""
            deadline = time.time() + timeout_seconds
            while time.time() < deadline:
                if _app_update_info is not None and _model_update_info is not None:
                    return
                time.sleep(0.1)

        def _build_app_banner_updates():
            """Return the three Gradio updates for the app-update banner."""
            info = _app_update_info
            if info is None or not info.update_available:
                return (gr.update(visible=False),) * 3
            changelog_lines = info.changelog.split("\n") if info.changelog else []
            if len(changelog_lines) > 20:
                changelog = "\n".join(changelog_lines[:20]) + "\n\n..."
            else:
                changelog = info.changelog or "No changelog available."
            date_str = f" ({info.published_date})" if info.published_date else ""
            banner_text = (
                f"### A new version is available: "
                f"v{info.current_version} \u2192 v{info.latest_version}{date_str}\n\n"
                f"**What's New:**\n\n{changelog}\n"
            )
            return (
                gr.update(value=banner_text, visible=True),
                gr.update(visible=True),
                gr.update(visible=True),
            )

        def _build_model_banner_updates():
            """Return the three Gradio updates for the model-update banner."""
            info = _model_update_info
            if info is None or not info.update_available:
                return (gr.update(visible=False),) * 3
            current_date = info.current_date or "unknown"
            latest_date = info.latest_date or "unknown"
            banner_text = (
                "### New PII model update available\n\n"
                f"Current: {current_date} \u2192 Latest: {latest_date}\n"
            )
            return (
                gr.update(value=banner_text, visible=True),
                gr.update(visible=True),
                gr.update(visible=True),
            )

        app.load(
            fn=check_and_update_banner,
            outputs=[update_banner, update_btn, later_btn, model_update_banner, model_update_btn, model_later_btn],
        )

    return app


def _start_update_check(app):
    """Start background update checks after a short delay."""
    threading.Thread(target=_check_app_update_background, daemon=True).start()
    threading.Thread(target=_check_model_update_background, daemon=True).start()


if __name__ == "__main__":
    print("=" * 50)
    print("  Privacy Filter - Local Interface")
    print("=" * 50)
    print()
    print("The model will be loaded the first time you use Detect.")
    print()

    app = create_ui()
    _start_update_check(app)
    app.queue()
    
    # Try port 7860, if busy try next ports
    port = 7860
    max_port = 7870
    while port <= max_port:
        try:
            app.launch(server_name="0.0.0.0", server_port=port, share=False)
            print(f"\nOpen http://localhost:{port}")
            break
        except OSError as e:
            if "already in use" in str(e) or "10048" in str(e):
                print(f"Port {port} in use, trying {port + 1}...")
                port += 1
            else:
                raise

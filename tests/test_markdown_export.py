"""Behaviour of the Markdown export module.

The endpoints exercise the module already through the integration tests, but
these are the unit-level checks: each branch does what the surrounding code
assumes it does, so a change to markitdown or mammoth cannot silently break
them.
"""

from __future__ import annotations

import io
import os
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

pytest.importorskip("markitdown")
pytest.importorskip("pymupdf")
pytest.importorskip("docx")

from server import markdown_export


@pytest.fixture(autouse=True)
def _reset_engine():
    """Drop the shared engine between tests so any patched settings apply."""
    markdown_export.clear_engine_cache()
    yield
    markdown_export.clear_engine_cache()


def _text_pdf(body: str) -> str:
    """A PDF with a real text layer, so markitdown's built-in reader can pick
    it up. Written to a temp file because MarkItDown.convert_stream still needs
    the extension to route to the PDF converter."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), body, fontsize=12)
    path = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
    doc.save(path)
    doc.close()
    return path


def _docx(placeholders: bool) -> str:
    """A DOCX with a heading, a paragraph and a small table. When
    ``placeholders`` is true the text carries the pipeline's redaction tokens
    that mammoth will otherwise escape as ``[NOMBRE\\_1]``."""
    import docx

    document = docx.Document()
    document.add_heading("Escritura anonimizada", level=1)
    if placeholders:
        document.add_paragraph(
            "Reunidos [NOMBRE_1] y [NOMBRE_2] con DNI [DNI_1]."
        )
    else:
        document.add_paragraph("Reunidos en Valencia.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Concepto"
    table.cell(0, 1).text = "Importe"
    table.cell(1, 0).text = "Notaria"
    table.cell(1, 1).text = "850"
    path = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
    document.save(path)
    return path


def test_pdf_with_text_layer_is_rendered_as_markdown() -> None:
    path = _text_pdf("Contrato firmado por [NOMBRE_1] con DNI [DNI_1]")
    try:
        markdown = markdown_export.to_markdown(path, ".pdf")
        assert "[NOMBRE_1]" in markdown
        assert "[DNI_1]" in markdown
    finally:
        os.unlink(path)


def test_docx_preserves_headings_and_tables() -> None:
    """The whole point of Markdown over plain text is that structure survives.

    The user pastes this into an LLM: knowing which lines are a heading and
    which cells belong to a table changes the model's answer.
    """
    path = _docx(placeholders=False)
    try:
        markdown = markdown_export.to_markdown(path, ".docx")
    finally:
        os.unlink(path)
    assert "# Escritura anonimizada" in markdown
    # mammoth uses GFM table syntax with |----| separator rows.
    assert "| Concepto | Importe |" in markdown
    assert "| Notaria | 850 |" in markdown


def test_placeholders_survive_the_mammoth_escape() -> None:
    """mammoth escapes ``_`` to ``\\_`` in Markdown, which would garble our tokens.

    The module post-processes that inside placeholder-shaped tokens so the
    user sees ``[NOMBRE_1]`` and not ``[NOMBRE\\_1]`` when they paste into an
    LLM.
    """
    path = _docx(placeholders=True)
    try:
        markdown = markdown_export.to_markdown(path, ".docx")
    finally:
        os.unlink(path)
    assert "[NOMBRE_1]" in markdown
    assert "[NOMBRE_2]" in markdown
    assert "[DNI_1]" in markdown
    # None of the tokens should have leaked with the escape still on.
    assert "\\_" not in markdown or "[NOMBRE\\_" not in markdown


def test_regular_content_underscore_escapes_are_left_alone() -> None:
    """The un-escape must only touch placeholder-shaped tokens."""
    text = r"palabra\_con\_guiones y un [NOMBRE\_1] tag"
    assert (
        markdown_export._unescape_placeholders(text)
        == r"palabra\_con\_guiones y un [NOMBRE_1] tag"
    )


def test_scanned_pdf_falls_back_to_the_pipeline_text() -> None:
    """A redacted scanned PDF has no text layer, so markitdown returns empty.

    The module must fall back to the already-redacted text the pipeline
    computed, not report an empty result to the user.
    """
    # An empty file mimics "markitdown could not read this". A real scanned
    # PDF would give a PDFSyntaxError of a different kind, but the branch we
    # care about is "markitdown returned nothing usable".
    bad = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
    Path(bad).write_bytes(b"")
    try:
        markdown = markdown_export.to_markdown(
            bad, ".pdf", fallback_text="ACTA firmada por [NOMBRE_1]"
        )
    finally:
        os.unlink(bad)
    assert "ACTA firmada por [NOMBRE_1]" in markdown


def test_unsupported_extension_returns_the_fallback() -> None:
    """The .txt / .md paths never produce a redacted file the module could
    open — they are returned as text upstream. The fallback text carries
    whatever the endpoint had."""
    assert markdown_export.to_markdown(
        "/dev/null", ".txt", fallback_text="soy texto plano [NOMBRE_1]"
    ) == "soy texto plano [NOMBRE_1]"


def test_no_fallback_and_no_conversion_returns_empty_string() -> None:
    """The endpoint decides what to do with an empty result (usually: don't
    even expose a download link). The module doesn't lie about producing
    content."""
    empty = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
    Path(empty).write_bytes(b"")
    try:
        assert markdown_export.to_markdown(empty, ".pdf") == ""
    finally:
        os.unlink(empty)

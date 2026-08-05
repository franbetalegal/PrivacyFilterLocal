"""Integration test: the ``also_markdown`` flag on ``/api/redact-file/apply``.

Only the /apply endpoint is exercised end-to-end (it does not need the model,
so it runs on the CI runner without loading anything heavy). The /redact-file
endpoint reuses the same helpers under the hood, so its behaviour is covered
by the unit tests in ``test_markdown_export.py`` plus the shared helper coverage
here.
"""

from __future__ import annotations

import io
import json
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

pytest.importorskip("fastapi")
from fastapi.testclient import TestClient

from server.main import app

try:
    client = TestClient(app)
except TypeError as exc:  # httpx/starlette version mismatch in this env
    pytest.skip(f"TestClient unavailable: {exc}", allow_module_level=True)


def _docx_bytes_with(name: str, dni: str) -> bytes:
    """A tiny DOCX we can compute redaction offsets over deterministically."""
    import docx

    document = docx.Document()
    document.add_heading("Escritura", level=1)
    document.add_paragraph(f"Reunidos {name} con DNI {dni}.")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _spans_for(name: str, dni: str) -> list[dict]:
    """The spans as the user would have curated after the review pass.

    The DOCX branch of the redaction only needs ``text`` and ``placeholder``,
    so offsets can be nominal — the runtime finds the substring itself.
    """
    return [
        {
            "start": 0,
            "end": len(name),
            "text": name,
            "placeholder": "[NOMBRE_1]",
            "label": "NOMBRE",
        },
        {
            "start": 0,
            "end": len(dni),
            "text": dni,
            "placeholder": "[DNI_1]",
            "label": "DNI",
        },
    ]


def test_apply_without_the_flag_gives_no_markdown_token():
    body = _docx_bytes_with("Juan Perez", "12345678Z")
    r = client.post(
        "/api/redact-file/apply",
        files={"file": ("escritura.docx", io.BytesIO(body),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"spans_json": json.dumps(_spans_for("Juan Perez", "12345678Z"))},
    )
    assert r.status_code == 200, r.text
    payload = r.json()
    assert payload["download_token"]
    assert payload.get("markdown_token") is None
    assert payload.get("markdown_name") is None


def test_apply_with_the_flag_produces_a_markdown_download():
    body = _docx_bytes_with("Juan Perez", "12345678Z")
    r = client.post(
        "/api/redact-file/apply",
        files={"file": ("escritura.docx", io.BytesIO(body),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={
            "spans_json": json.dumps(_spans_for("Juan Perez", "12345678Z")),
            "also_markdown": "true",
        },
    )
    assert r.status_code == 200, r.text
    payload = r.json()
    assert payload["download_token"]
    assert payload["markdown_token"]
    assert payload["markdown_name"] == "escritura_ANONIMIZED.md"

    # Download the .md and check the content: placeholders present, original
    # PII gone, heading preserved.
    md = client.get(f"/api/download/{payload['markdown_token']}")
    assert md.status_code == 200
    text = md.text
    assert "[NOMBRE_1]" in text
    assert "[DNI_1]" in text
    assert "Juan Perez" not in text
    assert "12345678Z" not in text
    assert "# Escritura" in text or "Escritura" in text

    # Second download of the same token is 404: single-use, same as the PDF.
    second = client.get(f"/api/download/{payload['markdown_token']}")
    assert second.status_code == 404


def test_apply_without_spans_still_honours_also_markdown():
    """A file with no detections still gets a .md if the user asked for it.

    Same output as feeding the original through markitdown directly, since no
    spans means the redacted copy is a byte-copy of the input.
    """
    body = _docx_bytes_with("Sin datos sensibles aqui", "0")
    r = client.post(
        "/api/redact-file/apply",
        files={"file": ("nota.docx", io.BytesIO(body),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"spans_json": "[]", "also_markdown": "true"},
    )
    assert r.status_code == 200, r.text
    payload = r.json()
    # No redactions were needed, so the output is a copy of the input, and its
    # markdown is that same document as markdown.
    assert payload["markdown_token"]
    md = client.get(f"/api/download/{payload['markdown_token']}")
    assert md.status_code == 200
    assert "Sin datos sensibles aqui" in md.text

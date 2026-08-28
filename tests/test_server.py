"""Backend API tests (model-free paths) using FastAPI's TestClient.

These cover routing, validation and serialization without loading the model.
If the installed httpx/starlette combination cannot build a TestClient, the
whole module is skipped (the assertions still run on a clean install where
requirements-server.txt pins a compatible stack).
"""

import io
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

pytest.importorskip("fastapi")
from fastapi.testclient import TestClient

from server.main import app

try:
    client = TestClient(app)
except TypeError as exc:  # httpx/starlette version mismatch in this env
    pytest.skip(f"TestClient unavailable: {exc}", allow_module_level=True)


def test_version():
    r = client.get("/api/version")
    assert r.status_code == 200
    assert "version" in r.json()


def test_health():
    r = client.get("/api/health")
    assert r.status_code == 200
    assert "model_loaded" in r.json()


def test_health_reports_every_detection_component():
    """The check that would have caught the three releases shipped without NER.

    An install missing the spaCy models detects no name written in caps, and
    from the outside that looks exactly like a document with no names in it.
    Health has to be able to say so.
    """
    components = client.get("/api/health").json()["components"]

    assert set(components) == {"opf", "ner", "ocr"}
    assert isinstance(components["ner"]["available"], bool)
    assert components["ner"]["models"], "no NER model is configured at all"
    for model in components["ner"]["models"]:
        assert model["name"]
        assert "present" in model


def test_health_is_not_ready_while_the_ner_layer_is_missing(monkeypatch):
    """``ready`` gates the whole UI, so it must not ignore the NER layer."""
    from server import inference

    real = inference.ner_status()
    monkeypatch.setattr(
        inference, "ner_status", lambda: {**real, "available": False}
    )

    assert client.get("/api/health").json()["ready"] is False


def test_startup_preflight_covers_every_component(monkeypatch):
    """The preflight is the only thing between a user and a half-configured
    detector, so it has to run all three stages, in order."""
    from server import inference

    ran: list[str] = []
    for stage in ("ensure_model_ready", "ensure_ner_models_ready", "refresh_ner_models"):
        monkeypatch.setattr(
            inference, stage, lambda s=stage: ran.append(s)
        )

    inference.run_preflight()

    assert ran == ["ensure_model_ready", "ensure_ner_models_ready", "refresh_ner_models"]


def test_health_carries_a_stable_instance_id():
    """The client tells a restart from a live server by this id changing.

    Without it, the page reloaded after an update would be served by the process
    that is still on its way out.
    """
    first = client.get("/api/health").json()
    second = client.get("/api/health").json()
    assert first["instance"]
    assert first["instance"] == second["instance"]


def test_redact_empty_text_is_noop():
    r = client.post("/api/redact", json={"text": "   "})
    assert r.status_code == 200
    body = r.json()
    assert body["empty"] is True
    assert body["detected_spans"] == []


def test_redact_file_rejects_unsupported_type():
    files = {"file": ("note.xyz", io.BytesIO(b"hello"), "application/octet-stream")}
    r = client.post("/api/redact-file", files=files)
    assert r.status_code == 400


def test_download_unknown_token_404():
    r = client.get("/api/download/does-not-exist")
    assert r.status_code == 404


def test_redact_empty_text_echoes_mode():
    """Mode is echoed even for empty input so the UI knows which preset ran."""
    r = client.post("/api/redact", json={"text": "  ", "mode": "conservative"})
    assert r.status_code == 200
    body = r.json()
    assert body["empty"] is True
    assert body["mode"] == "conservative"


def test_redact_empty_text_defaults_mode_to_balanced():
    r = client.post("/api/redact", json={"text": " "})
    assert r.status_code == 200
    assert r.json()["mode"] == "balanced"


def test_apply_rejects_unsupported_type():
    """/api/redact-file/apply only accepts PDF and DOCX."""
    files = {"file": ("note.txt", io.BytesIO(b"hello"), "text/plain")}
    r = client.post(
        "/api/redact-file/apply",
        files=files,
        data={"spans_json": "[]"},
    )
    assert r.status_code == 400


def test_apply_rejects_malformed_spans_json():
    files = {"file": ("doc.pdf", io.BytesIO(b"%PDF-1.4"), "application/pdf")}
    r = client.post(
        "/api/redact-file/apply",
        files=files,
        data={"spans_json": "{not-json"},
    )
    assert r.status_code == 400


def test_apply_docx_with_empty_span_list(tmp_path):
    """Empty span list must succeed and return a fresh download token."""
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    doc.add_paragraph("Some benign text without PII.")
    p = tmp_path / "clean.docx"
    doc.save(str(p))
    files = {"file": ("clean.docx", p.read_bytes(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = client.post(
        "/api/redact-file/apply",
        files=files,
        data={"spans_json": "[]"},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["applied_span_count"] == 0
    assert body["download_token"]
    assert body["download_name"] == "clean_ANONIMIZED.docx"


def test_apply_docx_with_curated_span_list(tmp_path):
    """A user-curated span list is applied verbatim (no re-detection)."""
    docx = pytest.importorskip("docx")
    doc = docx.Document()
    doc.add_paragraph("El testigo Juan García firmó el acta.")
    p = tmp_path / "in.docx"
    doc.save(str(p))
    files = {"file": ("in.docx", p.read_bytes(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    spans = [{
        "start": 11, "end": 22, "text": "Juan García",
        "placeholder": "[NOMBRE_1]", "label": "NOMBRE",
    }]
    r = client.post(
        "/api/redact-file/apply",
        files=files,
        data={"spans_json": __import__("json").dumps(spans)},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["applied_span_count"] == 1
    assert body["download_token"]

    # Fetch the redacted DOCX and verify the name is gone.
    download = client.get(f"/api/download/{body['download_token']}")
    assert download.status_code == 200
    out = tmp_path / "out.docx"
    out.write_bytes(download.content)
    from docx import Document
    text = "\n".join(x.text for x in Document(str(out)).paragraphs)
    assert "Juan García" not in text
    assert "[NOMBRE_1]" in text


def test_port_selection_skips_a_busy_port():
    """The port has to be chosen before uvicorn starts.

    uvicorn calls sys.exit(1) from Config.bind_socket when the bind fails, so
    the old retry loop around uvicorn.run (except OSError) never ran and the
    server died silently on a busy port.
    """
    import socket

    from server.main import _pick_port, _port_is_free

    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as taken:
        taken.bind(("127.0.0.1", 0))
        taken.listen(1)
        busy = taken.getsockname()[1]

        assert _port_is_free("127.0.0.1", busy) is False
        # Which port it lands on depends on what else the host has open; that it
        # moves past the busy one is the guarantee.
        chosen = _pick_port("127.0.0.1", busy)
        assert chosen is not None and chosen > busy


def test_port_selection_gives_up_instead_of_hanging():
    from server.main import _pick_port

    assert _pick_port("127.0.0.1", 7860, attempts=0) is None

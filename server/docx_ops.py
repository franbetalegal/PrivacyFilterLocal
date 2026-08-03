"""DOCX traversal helpers that reach every text container OOXML defines.

The default python-docx surface (``doc.paragraphs``) only exposes body-level
paragraphs, so text sitting in tables, headers, footers, text boxes or
comments is invisible to it. Legal writs routinely put PII in those exact
places (party names in footer letterheads, expedient numbers in header
tables, comments during review). This module walks every container and
yields ``Paragraph`` objects the existing run-splitter can rewrite.
"""
from __future__ import annotations

import logging
from typing import Iterator

logger = logging.getLogger("privacy_filter.docx_ops")


def _iter_table_paragraphs(table):
    """Yield paragraphs from every cell of ``table``, recursing into subtables."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_section_hf_paragraphs(section):
    """Yield paragraphs from every header/footer variant defined on a section."""
    for name in (
        "header", "footer",
        "first_page_header", "first_page_footer",
        "even_page_header", "even_page_footer",
    ):
        hf = getattr(section, name, None)
        if hf is None:
            continue
        try:
            yield from hf.paragraphs
            for table in hf.tables:
                yield from _iter_table_paragraphs(table)
        except Exception as exc:  # noqa: BLE001
            logger.debug("Skipping %s on section: %s", name, exc)


def _iter_textbox_paragraphs(doc):
    """Yield paragraphs stored inside ``w:txbxContent`` (text boxes)."""
    try:
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
    except ImportError:
        return

    parts = [doc.part]
    for section in doc.sections:
        for name in (
            "header", "footer",
            "first_page_header", "first_page_footer",
            "even_page_header", "even_page_footer",
        ):
            hf = getattr(section, name, None)
            if hf is not None and getattr(hf, "part", None) is not None:
                parts.append(hf.part)

    seen: set[int] = set()
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")
    for part in parts:
        if id(part) in seen:
            continue
        seen.add(id(part))
        root = getattr(part, "element", None)
        if root is None:
            continue
        for txbx in root.iter(txbx_tag):
            for p_elem in txbx.iter(p_tag):
                try:
                    yield Paragraph(p_elem, part)
                except Exception as exc:  # noqa: BLE001
                    logger.debug("Could not wrap textbox paragraph: %s", exc)


def _iter_comment_parts(doc):
    """Yield the comments-part element (``w:comments`` root), if present."""
    try:
        package = doc.part.package
    except AttributeError:
        return
    parts = getattr(package, "iter_parts", None)
    parts = parts() if callable(parts) else getattr(package, "parts", [])
    for part in parts:
        partname = getattr(part, "partname", None)
        if partname is None:
            continue
        if str(partname).endswith("comments.xml"):
            elem = getattr(part, "element", None)
            if elem is not None:
                yield part, elem


def _iter_comment_paragraphs(doc):
    """Yield paragraph-like wrappers for every ``w:p`` inside comments."""
    try:
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
    except ImportError:
        return
    p_tag = qn("w:p")
    for part, elem in _iter_comment_parts(doc):
        for p_elem in elem.iter(p_tag):
            try:
                yield Paragraph(p_elem, part)
            except Exception as exc:  # noqa: BLE001
                logger.debug("Could not wrap comment paragraph: %s", exc)


def iter_all_paragraphs(doc) -> Iterator:
    """Yield every paragraph the document exposes across every container.

    Order: body, body-tables, per-section headers/footers (+ their tables),
    text boxes, comments. The order is stable but not semantically meaningful
    for the pipeline — it flattens to text with ``\\n`` joins for detection.
    """
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        yield from _iter_section_hf_paragraphs(section)
    yield from _iter_textbox_paragraphs(doc)
    yield from _iter_comment_paragraphs(doc)


def extract_text(docx_path: str) -> str | None:
    """Extract text from every container joined by newlines. ``None`` on error."""
    try:
        from docx import Document
    except ImportError as exc:
        logger.error("python-docx is not installed: %s", exc)
        return None
    try:
        doc = Document(docx_path)
    except Exception as exc:  # noqa: BLE001
        logger.error("DOCX open failed: %s", exc)
        return None
    parts = [p.text for p in iter_all_paragraphs(doc)]
    return "\n".join(parts)
